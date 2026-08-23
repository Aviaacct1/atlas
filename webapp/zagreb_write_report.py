# Generate the Zagreb "Executive Summary & Assumptions" report as a Word doc, in the issued
# house style but with cleaner Avia formatting. Populated from the forecast + config assumptions.
# Usage: python zagreb_write_report.py [override_pack.json] [out.docx]   Author: Avia Solutions.
import json, sys, glob, os, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# The Zagreb folder resolves through avia_forecast/paths.py, like every other data
# location. This module and zagreb_write_excel.py each carried their own copy of the
# default, written differently (one with doubled separators), which is two owners for
# one constant. Author: Avia Solutions.
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from avia_forecast import paths

# Refusal rule (16 August 2026): a warned run never becomes a client artefact, whether
# reached over the service or run by hand. Clearing happens in
# config/export_watchpoints.yaml, deliberately and auditably; there is no override flag.
from avia_forecast.export_guard import refusal_message as _refusal
_r = _refusal("zagreb")
if _r:
    sys.exit(_r)

def _zag(sub):
    return os.path.join(paths.ZAGREB, sub)
INP=_zag("Inputs/03 Forecast Model 2026 007 202620515.xlsx")
OUT=sys.argv[2] if len(sys.argv)>2 else _zag("Zagreb Traffic Forecast - Executive Summary and Assumptions (Avia engine).docx")
OV=json.load(open(sys.argv[1])).get("overrides",{}) if (len(sys.argv)>1 and sys.argv[1] and os.path.exists(sys.argv[1])) else {}

# ---- headline forecast (compact reproduction of the engine demand chain) ----
import openpyxl
wb=openpyxl.load_workbook(INP,data_only=True)
gdp=wb["GDP"]; yrc={int(gdp.cell(row=5,column=c).value):c for c in range(1,120) if isinstance(gdp.cell(row=5,column=c).value,(int,float)) and 2000<=gdp.cell(row=5,column=c).value<=2055}
MR={"Croatia":14,"Schengen EEA":15,"Non-Schengen EEA":16,"North America":17,"Rest of World":18}
GDPidx={}
for m,r in MR.items():
    g={y:float(gdp.cell(row=r,column=c).value) for y,c in yrc.items() if isinstance(gdp.cell(row=r,column=c).value,(int,float))}
    last=max(g); rat=g[last]/g[last-1]
    for y in range(last+1,2049): g[y]=g[y-1]*rat
    GDPidx[m]={y:(g.get(y,g[min(g)]))/g[2025] for y in range(2019,2049)}
out=wb["Jul20 - YYYY"]; yc={int(out.cell(row=5,column=cc).value):cc for cc in range(1,60) if isinstance(out.cell(row=5,column=cc).value,(int,float)) and 2010<=out.cell(row=5,column=cc).value<=2055}
def srow(l):
    for r in range(9,150):
        for c in (2,3,4,5):
            v=out.cell(row=r,column=c).value
            if isinstance(v,str) and v.strip()==l: return {y:out.cell(row=r,column=cc).value for y,cc in yc.items() if isinstance(out.cell(row=r,column=cc).value,(int,float))}
    return {}
TOT=srow("Total passengers"); INTL=srow("Total international passengers"); DOM=srow("Total domestic passengers")
BT=TOT[2025]; BINT=INTL[2025]; BDOM=DOM[2025]
lcc=wb["LCC traffic"]; hdr={}
for r in range(1,12):
    ys=[(c,int(lcc.cell(row=r,column=c).value)) for c in range(1,60) if isinstance(lcc.cell(row=r,column=c).value,(int,float)) and 2015<=lcc.cell(row=r,column=c).value<=2055]
    if len(ys)>=5: hdr={y:c for c,y in ys}; break
LCCg={}
for r in range(1,45):
    if any(isinstance(lcc.cell(row=r,column=c).value,str) and 'total passengers' in str(lcc.cell(row=r,column=c).value).lower() for c in range(1,7)):
        LCCg={y:float(lcc.cell(row=r,column=cc).value) for y,cc in hdr.items() if isinstance(lcc.cell(row=r,column=cc).value,(int,float))}; break
LCC_BASE=LCCg.get(2025,1659200.0); NONLCC_BASE=BT-LCC_BASE
ELZ={"Croatia":0.8031,"Schengen EEA":0.9674,"Non-Schengen EEA":0.8414,"North America":0.8734,"Rest of World":0.9136}
EL=OV.get("el",ELZ)
# demand weights on MIDT true-O&D 2025 (long-haul demand routes via hub connections, so it
# exceeds the direct-seat share): Schengen 65.0% / non-Schengen short-haul 15.5% / long-haul 19.5% of intl
_oag={"Schengen EEA":0.650,"Non-Schengen EEA":0.155,"Long Haul":0.195}
W={"Croatia":BDOM,"Schengen EEA":BINT*_oag["Schengen EEA"],"Non-Schengen EEA":BINT*_oag["Non-Schengen EEA"],"North America":BINT*_oag["Long Haul"]*0.22,"Rest of World":BINT*_oag["Long Haul"]*0.78}
def organic(y):
    tw=sum(W.values()); return sum((W[m]/tw)*(GDPidx[m][y])**EL[m] for m in W)
LCC_NET={}
for y in range(2026,2049):
    if y in TOT and TOT[y]: LCC_NET[y]=TOT[y]-NONLCC_BASE*organic(y)
    else:
        gy=max(LCC_NET); LCC_NET[y]=LCC_NET[gy]*(LCC_NET[gy]/LCC_NET[gy-1])
if "lccSpot" in OV:
    sp={int(k):float(v) for k,v in OV["lccSpot"].items()}; xs=sorted(sp)
    def lccnet(y):
        if y<=xs[0]: return sp[xs[0]]
        if y>=xs[-1]: return sp[xs[-1]]
        for i in range(len(xs)-1):
            if xs[i]<=y<=xs[i+1]: a,b=xs[i],xs[i+1]; return sp[a]+(sp[b]-sp[a])*(y-a)/(b-a)
        return 0
else:
    def lccnet(y): return LCC_NET[y]
def total(y): return NONLCC_BASE*organic(y)+lccnet(y)
def dom(y): return BDOM*(GDPidx["Croatia"][y])**EL["Croatia"]*(1+OV.get("domUplift",0.0037))**(y-2025)
def m(v): return "%.1fm"%(v/1e6)
def cagr(a,b,n): return ((total(b)/(total(a) if a>2025 else BT))**(1/n)-1)*100 if True else 0
T27=total(2027); T47=total(2047); C2427=((total(2027)/TOT.get(2024,BT))**(1/3)-1)*100; C2747=((total(2047)/total(2027))**(1/20)-1)*100

# ---- build the document ----
TEAL=RGBColor(0x15,0x60,0x82); GREY=RGBColor(0x5f,0x71,0x85)
doc=Document()
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(10.5)
cp=doc.core_properties; cp.author="Avia Solutions"; cp.last_modified_by="Avia Solutions"; cp.title="Zagreb Airport Traffic Forecast"
def H(text,size=15,color=TEAL,space=10,after=6):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_before=Pt(space); p.paragraph_format.space_after=Pt(after); return p
def P(text,size=10.5,italic=False,color=None):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(size); r.italic=italic
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(6); return p
# Cover
for _ in range(6): doc.add_paragraph()
c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run("Zagreb Airport"); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=TEAL
c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run("Traffic Forecast 2026-2048"); r.font.size=Pt(18); r.font.color.rgb=GREY
c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run("Executive Summary and Assumptions"); r.font.size=Pt(14); r.font.color.rgb=GREY
for _ in range(2): doc.add_paragraph()
c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run(datetime.date.today().strftime("%B %Y")); r.font.size=Pt(12)
c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run("PRIVATE & CONFIDENTIAL"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=GREY
c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run("Produced by the Avia global forecast engine (Zagreb configured instance)"); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GREY
doc.add_page_break()
# Disclaimer
H("Disclaimer of Liability",12)
P("This publication provides general information and should not be used or taken as business, financial, tax, accounting, legal or other advice, or relied upon in substitution for the exercise of your independent judgment. For your specific situation or where otherwise required, expert advice should be sought. Although Avia Solutions Limited or any of its affiliates (together, “Avia”) believes that the information contained in this publication has been obtained from and is based upon sources Avia believes to be reliable, Avia does not guarantee its accuracy and it may be incomplete or condensed. Avia makes no representation or warranties of any kind whatsoever in respect of such information. Avia accepts no liability of any kind for loss arising from the use of the material presented in this publication.",9)
H("Copyright",12)
P("Copyright © %d Avia Solutions Limited. All rights reserved."%datetime.date.today().year,9)
doc.add_page_break()
# Executive summary
H("Executive Summary",16)
P("This traffic forecast for Zagreb Airport (ZAG) is the latest in a series of annual updates to the mid-term forecast provided by MZLZ and reviewed by Avia Solutions. It carries forward the assumptions, carrier intentions and market intelligence of successive forecasts, together with the recent stated intentions and actions of Croatia Airlines and Ryanair.")
P("Zagreb has recovered strongly from the COVID-19 pandemic, with all segments except long haul above 2019 levels. Engagement with Ryanair continues to shape short to medium term growth and the structure of traffic, while Croatia Airlines refocuses on Zagreb through a full renewal of its fleet to Airbus A220 aircraft. Wizz Air is assumed to enter in summer 2026, moving to year-round service from 2027.")
P("On this basis total passenger throughput is forecast at %s in 2027, growing at %.1f%% per year to %s by 2047. The near term is anchored to forward schedules and carrier intentions; from 2029 growth is driven by GDP per market with continued low-cost stimulation. International traffic grows faster than domestic on the back of the A220 fleet and the Wizz Air entry."%(m(T27),C2747,m(T47)))
P("This report sets out the executive summary and the assumptions underpinning the forecast. Detailed passenger, movement and tonnage tables are provided in the accompanying Excel model.")
# Assumptions
doc.add_page_break()
H("1  ZAG Traffic Forecast Assumptions",16)
ASSUMP=[
("Short-term (bottom-up) forecast","The 2026-2028 forecast is built from current services, forward schedules (OAG), market intelligence (Routes meetings, MZLZ insight) and actual performance year to date. It blends econometric growth with aviation-industry factors. Croatia Airlines returns to 2019 passenger levels by 2028 and differentiates by product rather than competing with Ryanair on price or route."),
("Historic passenger data","Historic traffic 2019 to date is sourced from MZLZ tower logs and Sabre MIDT (2010-2025), treated as reliable for forecasting once pandemic and one-off events are excluded."),
("Passenger itineraries","MIDT (2010-2025) is used to identify non-stop, connecting and hubbing passengers and to categorise traffic into domestic O&D, international O&D and transfer types."),
("Key markets","MIDT analysis identifies the O&D markets used for the econometric work: Domestic (Croatia), Schengen EEA, Non-Schengen EEA, North America and Rest of World, aligned to the Croatian border regime."),
("Domestic O&D methodology","Croatian domestic passengers are regressed on Croatian GDP using 2013-2024 MIDT data, excluding the pandemic. Applied elasticity %.2f."%ELZ["Croatia"]),
("International O&D methodology","Each international market is regressed on its GDP series using 2013-2019 passengers. Applied elasticities: Schengen EEA %.2f, Non-Schengen EEA %.2f, North America %.2f, Rest of World %.2f."%(ELZ["Schengen EEA"],ELZ["Non-Schengen EEA"],ELZ["North America"],ELZ["Rest of World"])),
("GA / executive / other passengers","Assumed a fixed proportion of international and domestic commercial passengers, on the 2024-2025 average."),
("Transit passengers","Assumed a fixed proportion of commercial passengers on the 2024-2025 average, growing with total traffic."),
("GDP sources","Historic and forecast GDP are from OEF and IMF (composite). GDP underpins growth to 2028 alongside low-cost engagement; beyond 2028 GDP is the main driver for Croatia Airlines and other carriers, with continued LCC stimulation. Long-term GDP uses Euro area, Emerging and Developing Europe and World growth for EEA, Non-EEA and long-haul respectively."),
("Transfer (hubbing) passengers","Transfer passengers grow in line with total traffic on stronger European and long-haul connectivity. Low-cost traffic contributes self-connect rather than pure transfer."),
("Infants (children up to 2 years)","A constant percentage of domestic and international passengers on the 2024-2025 average."),
("EEA / Schengen split, final destination","Passenger splits by final destination are derived from 2019 and 2024-2025 MIDT. UK, Norway and Switzerland are outside EEA per the Croatian border regime."),
("EEA / Schengen split, aircraft destination","A direct output of the bottom-up model, with no assumed change in EEA or Schengen membership over the horizon (Croatia within Schengen from 2023)."),
("Scheduled vs non-scheduled","The ratio follows the consistent 2019 and 2023-2025 pattern. Low-cost service, particularly Ryanair, affects non-scheduled (charter) traffic."),
("Airport infrastructure and ATC","Assumed able to accommodate the forecast without constraining growth over the period."),
("Air service agreements","Existing bilateral agreements are not a material constraint to long-term development."),
("Air fares and oil prices","Not modelled separately; their correlation with GDP means the historic relationship is captured in the GDP-traffic regressions."),
("Airport charges","Tariff changes are not expected to materially affect full-service or non-scheduled demand. The five-year incentive scheme is central to low-cost service; a further round from 2026 is assumed, enabling continued Ryanair growth and Wizz Air entry (summer 2026, two aircraft, year-round from 2027)."),
("Long-haul forecast","The bottom-up long-haul set assumes DXB, YYZ, TLV, DOH and SEL. Recovery is gradual, being the last segment to return and exposed to geopolitical risk; Korea has resumed with low-cost T'Way."),
("Commercial ATM forecast","Average aircraft size with planning load factors by carrier type and segment gives ATMs for international, domestic, Croatia Airlines, LCC and other movements."),
("GA / executive / other ATMs","Assumed to grow in line with commercial ATMs."),
("Landed tonnage","Croatia Airlines operates 22t-80t aircraft, transitioning from Dash 8-400 and A319/320 to A220-100/300 with an ATR-72 proxy for wet-leased turboprops. Ryanair uses 737-800/MAX and Lauda A320."),
("Departing vs total","Departing data is total divided by two; departing non-scheduled commercial passengers include transit."),
("Monthly splits","Monthly traffic follows historical seasonality (summer peak, winter trough), with additional low-cost capacity adjusted month by month to promote winter traffic."),
("Non-commercial ATM and MTOW","The ratio between MTOW categories of non-commercial ATMs is held constant over the forecast."),
]
for i,(t,txt) in enumerate(ASSUMP,1):
    p=doc.add_paragraph(); r=p.add_run("Assumption %d  %s"%(i,t)); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=TEAL
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    P(txt)
# Fleet table
doc.add_page_break()
H("2  Croatia Airlines Fleet Renewal Assumptions",16)
P("Croatia Airlines renews its fleet from Airbus A319/A320 and Dash 8-400 to the Airbus A220 family (A220-100 and A220-300), with an ATR-72 proxy for wet-leased turboprops. Year-end fleet by type:")
try:
    fl=json.load(open(_zag("work/zagreb_fleet.json")))["by_year"]
except Exception:
    fl={}
years=[y for y in ["2024","2025","2026","2027"] if y in fl]
tbl=doc.add_table(rows=1,cols=1+len(years)); tbl.style="Light Grid Accent 1"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr=tbl.rows[0].cells; hdr[0].text="Aircraft"
for j,y in enumerate(years): hdr[j+1].text=y
for tp,label in [("319","A319"),("320","A320"),("DH4","Dash 8-400"),("221","A220-100"),("223","A220-300")]:
    row=tbl.add_row().cells; row[0].text=label
    for j,y in enumerate(years): row[j+1].text=str(fl[y]["fleet"].get(tp,0))
row=tbl.add_row().cells; row[0].text="Total"
for j,y in enumerate(years): row[j+1].text=str(fl[y]["total_ac"])
P("The first A220-100 enters limited service in late 2025; three ATR-72 are wet-leased to replace the remaining Dash 8-400s. Croatia Airlines' long-term capacity is higher than previously anticipated on a greater number of the larger A220-300.",9,italic=True)
# Summary tables
doc.add_page_break()
H("3  Forecast Summary",16)
spots=[2026,2027,2032,2037,2042,2047]
def add_tbl(title,rows):
    P(title,11); t=doc.add_table(rows=1,cols=1+len(spots)); t.style="Light Grid Accent 1"
    h=t.rows[0].cells; h[0].text="'000 passengers"
    for j,y in enumerate(spots): h[j+1].text=str(y)
    for label,fn in rows:
        rc=t.add_row().cells; rc[0].text=label
        for j,y in enumerate(spots): rc[j+1].text="{:,}".format(round(fn(y)/1000))
    doc.add_paragraph()
add_tbl("Passengers by type",[("Total passengers",total),("International",lambda y: total(y)-dom(y)-0),("Domestic",dom)])
P("Source: Sabre GDD, OEF/IMF, OAG; Avia Solutions analysis. Figures produced by the Avia global forecast engine (Zagreb configured instance) and subject to analyst review before issue.",9,italic=True,color=GREY)
doc.save(OUT)
print("saved report:",OUT)
